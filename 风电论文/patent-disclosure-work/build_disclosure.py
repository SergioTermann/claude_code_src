from __future__ import annotations

import json
import math
import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, FancyArrowPatch, FancyBboxPatch, Polygon, Rectangle, Wedge

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "patent-disclosure-work"
FIG = OUT / "cn-figures"
FIG.mkdir(parents=True, exist_ok=True)
SKILL_SCRIPTS = Path("/Users/zinger/.codex/skills/nature-paper-to-patent/scripts")
sys.path.insert(0, str(SKILL_SCRIPTS))

from math_to_omml import latex_to_omml

plt.rcParams["font.sans-serif"] = [
    "Arial Unicode MS",
    "Hiragino Sans GB",
    "Heiti TC",
    "Songti SC",
    "DejaVu Sans",
]
plt.rcParams["axes.unicode_minus"] = False

RICH = {
    "blue": ("#DDEBFA", "#2F6FA8"),
    "teal": ("#DDF4ED", "#2C9C88"),
    "green": ("#E4F4D8", "#5E9F3D"),
    "amber": ("#FFF0C7", "#C28A14"),
    "orange": ("#FFE2CC", "#D46A22"),
    "red": ("#F7D8D8", "#B94C4C"),
    "purple": ("#E8DEF8", "#7B5BB4"),
    "magenta": ("#F7D7EA", "#B44786"),
    "cyan": ("#D8F1F7", "#3A91A8"),
    "gray": ("#EEF1F4", "#59616A"),
}

NODE_COLORS = [
    RICH["blue"],
    RICH["orange"],
    RICH["teal"],
    RICH["purple"],
    RICH["magenta"],
    RICH["green"],
    RICH["amber"],
    RICH["cyan"],
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def setup_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Title"]:
        styles[name].font.name = "黑体"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(18)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(14)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(12)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def add_math(doc: Document, latex: str, number: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_para = latex_to_omml(latex)
    p._p.append(math_para[0])
    r = p.add_run(f"    （{number}）")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        r = p.add_run("• " + item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        r = p.add_run(f"{i}. {item}")
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)


def draw_box(ax, xy, wh, label, fill="#FFFFFF", edge="#333333", size=10, bold=False):
    x, y = xy
    w, h = wh
    ax.add_patch(Rectangle((x, y), w, h, facecolor=fill, edgecolor=edge, linewidth=1.2))
    ax.text(
        x + w / 2,
        y + h / 2,
        label,
        ha="center",
        va="center",
        fontsize=size,
        fontweight="bold" if bold else "normal",
        wrap=True,
    )


def draw_circle(ax, xy, radius, label, fill="#FFFFFF", edge="#333333", size=9, bold=False):
    ax.add_patch(Circle(xy, radius, facecolor=fill, edgecolor=edge, linewidth=1.2))
    ax.text(
        xy[0],
        xy[1],
        label,
        ha="center",
        va="center",
        fontsize=size,
        fontweight="bold" if bold else "normal",
        wrap=True,
    )


def draw_pill(ax, xy, wh, label, fill="#FFFFFF", edge="#333333", size=8, bold=False):
    x, y = xy
    w, h = wh
    ax.add_patch(
        FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle=f"round,pad=0.006,rounding_size={h / 2}",
            facecolor=fill,
            edgecolor=edge,
            linewidth=1.2,
        )
    )
    ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=size, fontweight="bold" if bold else "normal")


def curved_arrow(ax, start, end, rad=0.2, color="#333333", dashed=False):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.2,
            linestyle="--" if dashed else "-",
            color=color,
            connectionstyle=f"arc3,rad={rad}",
            shrinkA=4,
            shrinkB=4,
        )
    )


def arrow(ax, start, end):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.2,
            color="#333333",
            shrinkA=4,
            shrinkB=4,
        )
    )


def dashed_arrow(ax, start, end, rad=0.0):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.2,
            linestyle="--",
            color="#555555",
            connectionstyle=f"arc3,rad={rad}",
            shrinkA=4,
            shrinkB=4,
        )
    )


def poly_arrow(ax, points, dashed=False):
    for start, end in zip(points[:-2], points[1:-1]):
        ax.plot(
            [start[0], end[0]],
            [start[1], end[1]],
            color="#333333",
            linewidth=1.2,
            linestyle="--" if dashed else "-",
            clip_on=False,
        )
    style = "--" if dashed else "-"
    ax.add_patch(
        FancyArrowPatch(
            points[-2],
            points[-1],
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.2,
            linestyle=style,
            color="#333333",
            shrinkA=0,
            shrinkB=4,
            clip_on=False,
        )
    )


def save_fig(fig, name: str) -> Path:
    path = FIG / f"{name}.png"
    fig.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0.08)
    plt.close(fig)
    return path


def fig_main_flow() -> Path:
    fig, ax = plt.subplots(figsize=(12.4, 7.0))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.96, "图1  本地化风电运维证据约束LLM系统架构", ha="center", fontsize=14, fontweight="bold")

    # Central evidence bus, matching the paper's "local evidence supplies facts" logic.
    draw_pill(ax, (0.18, 0.455), (0.64, 0.10), "本地证据总线：来源路径 | 机型范围 | 图谱邻居 | 机理链路 | 故障状态", fill=RICH["red"][0], edge=RICH["red"][1], size=10, bold=True)

    data_nodes = [
        ("本地手册\n结构/流程", 0.14, 0.78),
        ("故障码表\n代码/机型", 0.32, 0.82),
        ("案例报告\n现象/原因", 0.50, 0.80),
        ("维修工单\n反馈/闭环", 0.68, 0.82),
        ("问答轨迹\n多轮状态", 0.84, 0.76),
    ]
    for i, (label, x, y) in enumerate(data_nodes):
        fill, edge = NODE_COLORS[i]
        draw_circle(ax, (x, y), 0.058, label, fill=fill, edge=edge, size=7.2, bold=True)
        curved_arrow(ax, (x, y - 0.058), (x, 0.555), rad=0.0)
    ax.text(0.06, 0.84, "数据输入", fontsize=11, fontweight="bold", color=RICH["blue"][1])

    knowledge_nodes = [
        ("字段归一化\n风场/机型/部件", 0.19, 0.31),
        ("故障条目\n来源+原文", 0.36, 0.25),
        ("知识图谱\n22,680/79,474", 0.53, 0.29),
        ("机理模板\n12类", 0.68, 0.25),
        ("Wiki/索引\n可读可检索", 0.83, 0.31),
    ]
    for i, (label, x, y) in enumerate(knowledge_nodes):
        fill, edge = NODE_COLORS[i + 3]
        draw_circle(ax, (x, y), 0.055, label, fill=fill, edge=edge, size=7.0, bold=True)
        curved_arrow(ax, (x, 0.455), (x, y + 0.055), rad=0.0)
    ax.text(0.06, 0.30, "知识资产", fontsize=11, fontweight="bold", color=RICH["teal"][1])

    reasoning = [
        ("意图路由", 0.18),
        ("范围绑定", 0.33),
        ("图增强检索", 0.50),
        ("策略链控制", 0.67),
        ("受约束生成", 0.82),
    ]
    for i, (label, x) in enumerate(reasoning):
        fill, edge = NODE_COLORS[(i + 5) % len(NODE_COLORS)]
        draw_circle(ax, (x, 0.12), 0.045, label, fill=fill, edge=edge, size=7.3, bold=True)
        if i:
            arrow(ax, (reasoning[i - 1][1] + 0.045, 0.12), (x - 0.045, 0.12))
    curved_arrow(ax, (0.50, 0.455), (0.50, 0.165), rad=0.0)
    draw_pill(ax, (0.20, 0.02), (0.60, 0.055), "输出：故障含义 | 适用范围 | 可能机理 | 验证动作 | 接受准则 | 证据路径", fill=RICH["purple"][0], edge=RICH["purple"][1], size=9, bold=True)
    curved_arrow(ax, (0.82, 0.12), (0.80, 0.075), rad=0.0)
    ax.text(0.50, 0.60, "LLM只组织证据包，不作为独立事实来源", ha="center", fontsize=9.2, color="#AA4444", fontweight="bold")
    return save_fig(fig, "图1_系统架构")


def fig_architecture() -> Path:
    fig, ax = plt.subplots(figsize=(12.4, 7.0))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.96, "图2  本地运维知识构建流程", ha="center", fontsize=14, fontweight="bold")
    ax.text(0.14, 0.86, "异构本地资料", ha="center", fontsize=11, fontweight="bold", color=RICH["blue"][1])
    ax.text(0.50, 0.86, "确定性解析流水", ha="center", fontsize=11, fontweight="bold", color=RICH["orange"][1])
    ax.text(0.82, 0.86, "可回溯知识资产", ha="center", fontsize=11, fontweight="bold", color=RICH["purple"][1])

    docs = [
        ("手册\n结构/流程", 0.12, 0.72, 0.058),
        ("故障码表\n代码/机型", 0.07, 0.60, 0.052),
        ("工单\n动作/反馈", 0.17, 0.57, 0.052),
        ("案例报告\n现象/原因", 0.12, 0.45, 0.058),
    ]
    for i, (label, x, y, r) in enumerate(docs):
        fill, edge = NODE_COLORS[i]
        draw_circle(ax, (x, y), r, label, fill=fill, edge=edge, size=7.2, bold=True)

    stage_x = [0.34, 0.45, 0.56, 0.67]
    stages = ["登记\n路径/版本", "解析\n文本/表格", "识别\n领域字段", "归一化\n别名/范围"]
    for i, (x, label) in enumerate(zip(stage_x, stages)):
        fill, edge = NODE_COLORS[i + 4]
        draw_circle(ax, (x, 0.63), 0.052, label, fill=fill, edge=edge, size=7.2, bold=True)
        if i:
            arrow(ax, (stage_x[i - 1] + 0.052, 0.63), (x - 0.052, 0.63))
    draw_pill(ax, (0.32, 0.44), (0.39, 0.07), "证据化转换核心：字段规范化 + 来源路径 + 原文片段 + 图谱关系", fill=RICH["red"][0], edge=RICH["red"][1], size=8.4, bold=True)
    for _, x, y, r in docs:
        curved_arrow(ax, (x + r, y), (0.29, 0.63), rad=0.10)
    curved_arrow(ax, (0.56, 0.58), (0.52, 0.51), rad=0.0)

    products = [
        ("故障索引\n4,849故障码节点", 0.80, 0.70, RICH["amber"][0], RICH["amber"][1]),
        ("知识图谱\n22,680节点\n79,474关系", 0.91, 0.59, RICH["teal"][0], RICH["teal"][1]),
        ("机理模板\n12类模板", 0.80, 0.45, RICH["purple"][0], RICH["purple"][1]),
        ("Wiki页面\n可读可检索", 0.91, 0.31, RICH["cyan"][0], RICH["cyan"][1]),
    ]
    for i, (label, x, y, fill, edge) in enumerate(products):
        draw_circle(ax, (x, y), 0.060, label, fill=fill, edge=edge, size=6.8, bold=True)
        curved_arrow(ax, (0.71, 0.475), (x - 0.06, y), rad=-0.10 + i * 0.06)
    for i in range(len(products) - 1):
        x1, y1 = products[i][1], products[i][2]
        x2, y2 = products[i + 1][1], products[i + 1][2]
        ax.plot([x1, x2], [y1, y2], color="#D0D0D0", linewidth=1.0, linestyle=":")
    schema = [
        ("故障条目 f_i", "故障码c_i | 机型m_i | 品牌b_i | 系统s_i | 部件p_i | 原因r_i | 动作a_i | 证据e_i"),
        ("机理模板 g_j", "症状u_j | 因果因素v_j | 优先检查q_j | 接受准则k_j"),
    ]
    for i, (name, fields) in enumerate(schema):
        y = 0.16 - i * 0.075
        fill, edge = (RICH["amber"] if i == 0 else RICH["purple"])
        draw_box(ax, (0.11, y), (0.16, 0.055), name, fill=fill, edge=edge, size=8, bold=True)
        draw_box(ax, (0.285, y), (0.60, 0.055), fields, fill="#FFFFFF", edge=edge, size=7.5)
    ax.text(0.5, 0.045, "流程目标：从非结构化本地资料生成可检索、可回溯、可供策略链调用的结构化证据资产。", ha="center", fontsize=9)
    return save_fig(fig, "图2_知识构建流程")


def fig_mechanism() -> Path:
    fig, ax = plt.subplots(figsize=(12.4, 7.0))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.96, "图4  机理增强推理图与假设鉴别闭环", ha="center", fontsize=14, fontweight="bold")
    center = (0.42, 0.55)
    draw_circle(ax, center, 0.075, "故障案例\nx", fill=RICH["blue"][0], edge=RICH["blue"][1], size=8.5, bold=True)
    rings = [(0.16, "#F4EAF5"), (0.27, "#EAF6F1")]
    for r, col in rings:
        ax.add_patch(Circle(center, r, facecolor="none", edgecolor=col, linewidth=8, alpha=0.9))
    nodes = [
        ("机理原型\na", 105, 0.18),
        ("机理层\nl", 60, 0.22),
        ("传播步骤\np", 18, 0.24),
        ("失效模式\nz", -25, 0.22),
        ("可观测量\no", -72, 0.19),
        ("验证试验\nt", -128, 0.20),
        ("控制屏障\nb", 170, 0.23),
    ]
    prev = center
    positions = []
    for i, (label, deg, radius) in enumerate(nodes):
        ang = math.radians(deg)
        pos = (center[0] + math.cos(ang) * radius, center[1] + math.sin(ang) * radius)
        positions.append(pos)
        fill, edge = NODE_COLORS[i + 1]
        draw_circle(ax, pos, 0.050, label, fill=fill, edge=edge, size=7.3, bold=True)
        curved_arrow(ax, prev, pos, rad=0.08)
        prev = pos
    curved_arrow(ax, positions[-1], center, rad=0.25, dashed=True)
    examples = [
        ("液压流量受限", 0.11, 0.79),
        ("机械接触疲劳", 0.28, 0.83),
        ("电气热-绝缘应力", 0.51, 0.83),
        ("传感采集反馈", 0.66, 0.76),
        ("通信时序一致", 0.69, 0.31),
        ("保护链边界", 0.17, 0.25),
    ]
    for txt, x, y in examples:
        fill, edge = NODE_COLORS[int((x * 100) % len(NODE_COLORS))]
        draw_pill(ax, (x, y), (0.14, 0.045), txt, fill="#FFFFFF", edge=edge, size=7)
    # Hypothesis discrimination branch as a separate arc.
    draw_pill(ax, (0.66, 0.54), (0.25, 0.06), "诊断假设 h", fill=RICH["orange"][0], edge=RICH["orange"][1], size=8.5, bold=True)
    hypo_nodes = [
        ("鉴别证据\ne_d", 0.73, 0.43),
        ("反事实测试\nt_c", 0.83, 0.35),
        ("判定规则\nr_d", 0.88, 0.23),
    ]
    start = (0.79, 0.54)
    for label, x, y in hypo_nodes:
        fill, edge = NODE_COLORS[(int(x * 100) + 2) % len(NODE_COLORS)]
        draw_circle(ax, (x, y), 0.052, label, fill=fill, edge=edge, size=7.2, bold=True)
        curved_arrow(ax, start, (x, y), rad=-0.1)
        start = (x, y)
    stats = [
        ("成对机理竞争\n18例", 0.61, 0.14, RICH["cyan"][0], RICH["cyan"][1]),
        ("单机理反事实\n15例", 0.75, 0.10, RICH["green"][0], RICH["green"][1]),
        ("假设鉴别覆盖\n33/33", 0.89, 0.14, RICH["magenta"][0], RICH["magenta"][1]),
    ]
    for label, x, y, fill, edge in stats:
        draw_circle(ax, (x, y), 0.055, label, fill=fill, edge=edge, size=6.8, bold=True)
    curved_arrow(ax, center, (0.66, 0.57), rad=-0.18)
    ax.text(0.5, 0.045, "径向机理网络表达论文中的闭环链路 x→a→l→p→z→o→t→b；右侧假设分支用于排除传感误差、边界条件或伪原因。", ha="center", fontsize=9)
    return save_fig(fig, "图4_机理增强推理图")


def fig_scope_strategy() -> Path:
    fig, ax = plt.subplots(figsize=(12.4, 6.6))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.95, "图3  模型范围绑定、证据包生成与策略链交互", ha="center", fontsize=14, fontweight="bold")

    # Radar-like scope binding view.
    center = (0.31, 0.58)
    labels = [
        ("C_i\n故障码", 90, 0.17),
        ("M_i\n机型", 30, 0.15),
        ("B_i\n品牌", -30, 0.14),
        ("S_i\n系统/部件", -90, 0.16),
        ("H_i\n会话", -150, 0.13),
        ("E_i\n证据", 150, 0.15),
    ]
    for r in [0.06, 0.11, 0.16]:
        ax.add_patch(Circle(center, r, facecolor="none", edgecolor="#D7E7DD", linewidth=1.0))
    pts = []
    for label, deg, radius in labels:
        ang = math.radians(deg)
        x = center[0] + math.cos(ang) * 0.20
        y = center[1] + math.sin(ang) * 0.20
        vx = center[0] + math.cos(ang) * radius
        vy = center[1] + math.sin(ang) * radius
        pts.append((vx, vy))
        ax.plot([center[0], x], [center[1], y], color="#C8D8D0", linewidth=0.8)
        ax.text(x, y, label, ha="center", va="center", fontsize=7.5, fontweight="bold")
    ax.add_patch(Polygon(pts, closed=True, facecolor=RICH["green"][0], edgecolor=RICH["green"][1], alpha=0.72, linewidth=1.5))
    draw_circle(ax, center, 0.045, "score\n(f_i,q)", fill=RICH["amber"][0], edge=RICH["amber"][1], size=7.5, bold=True)
    draw_box(ax, (0.05, 0.50), (0.15, 0.15), "维修问题q\n故障码/现象\n机型线索/短反馈", fill=RICH["blue"][0], edge=RICH["blue"][1], size=8, bold=True)
    arrow(ax, (0.20, 0.575), (0.265, 0.575))

    # Evidence packet as a ring, not boxes.
    packet_center = (0.62, 0.50)
    draw_circle(ax, packet_center, 0.12, "证据包\n适用范围\n事实/机理/来源", fill=RICH["red"][0], edge=RICH["red"][1], size=8.5, bold=True)
    fields = [
        ("结构化\n故障条目", 0.47, 0.64),
        ("图谱\n邻居", 0.64, 0.70),
        ("机理\n链路", 0.78, 0.55),
        ("来源\n路径", 0.70, 0.33),
        ("验证动作\n接受准则", 0.50, 0.32),
    ]
    for i, (label, x, y) in enumerate(fields):
        fill, edge = NODE_COLORS[i + 2]
        draw_circle(ax, (x, y), 0.050, label, fill=fill, edge=edge, size=6.8, bold=True)
        curved_arrow(ax, (x, y), packet_center, rad=0.08, color="#555555")
    curved_arrow(ax, (center[0] + 0.16, center[1]), (packet_center[0] - 0.12, packet_center[1] + 0.02), rad=-0.08)

    draw_pill(ax, (0.80, 0.40), (0.16, 0.20), "策略链输出\n故障含义\n可能机理\n一个现场验证动作\n接受准则\n证据路径", fill=RICH["purple"][0], edge=RICH["purple"][1], size=8, bold=True)
    arrow(ax, (packet_center[0] + 0.12, packet_center[1]), (0.80, 0.50))
    dashed_arrow(ax, (0.88, 0.40), (0.12, 0.50), rad=-0.28)
    ax.text(0.50, 0.08, "短反馈不重新开题，而是作为上一现场验证动作的观测结果回写到会话状态 H_i。", ha="center", fontsize=9)
    return save_fig(fig, "图3_范围绑定策略链")


def fig_metrics() -> Path:
    eval_path = ROOT.parent / "generated-knowledge" / "windrise-mechanism-graph-evaluation.json"
    data = json.loads(eval_path.read_text(encoding="utf-8"))
    fig, axes = plt.subplots(2, 2, figsize=(11.8, 7.2))
    fig.suptitle("图5  知识资产规模、机理图谱组成与覆盖效果", fontsize=14, fontweight="bold")
    scale_names = ["原始故障\n记录", "图谱节点", "图谱关系", "故障码\n节点", "来源文档\n节点"]
    scale_vals = [11865, 22680, 79474, 4849, 8295]
    ax00, ax01, ax10, ax11 = axes.ravel()

    ax00.set_title("本地知识资产规模")
    ax00.axis("off")
    bubble_pos = [(0.18, 0.62), (0.42, 0.68), (0.66, 0.58), (0.36, 0.32), (0.62, 0.30)]
    colors = ["#3A6EA5", "#4E9A80", "#C18F2C", "#8C6BB1", "#C05A5A"]
    max_v = max(scale_vals)
    for (x, y), name, val, color in zip(bubble_pos, scale_names, scale_vals, colors):
        radius = 0.055 + 0.075 * math.sqrt(val / max_v)
        ax00.add_patch(Circle((x, y), radius, facecolor=color, edgecolor="#333333", alpha=0.22, linewidth=1.3, transform=ax00.transAxes))
        ax00.text(x, y + 0.012, name, ha="center", va="center", fontsize=8.2, fontweight="bold", transform=ax00.transAxes)
        ax00.text(x, y - 0.060, f"{val:,}", ha="center", va="center", fontsize=8, transform=ax00.transAxes)
    ax00.text(0.5, 0.08, "覆盖13品牌、66机型、29风场、18系统、12类机理模板", ha="center", fontsize=8.2, transform=ax00.transAxes)

    ax01.set_title("闭环覆盖提升")
    ax01.axis("off")
    gauges = [
        ("传统画像\n完整率", data["baseline"]["complete_profile_rate"], "#999999"),
        ("机理画像\n完整率", data["mechanism"]["profile_complete_rate"], "#4E9A80"),
        ("验证闭环\n覆盖率", data["mechanism"]["validation_closure_rate"], "#3A6EA5"),
        ("预防闭环\n覆盖率", data["mechanism"]["prevention_closure_rate"], "#C18F2C"),
        ("假设鉴别\n覆盖率", data["mechanism"]["discrimination_coverage_rate"], "#8C6BB1"),
    ]
    for i, (label, val, color) in enumerate(gauges):
        x = 0.12 + i * 0.19
        y = 0.56
        ax01.add_patch(Wedge((x, y), 0.085, 180, 0, width=0.025, facecolor="#E8E8E8", edgecolor="none", transform=ax01.transAxes))
        ax01.add_patch(Wedge((x, y), 0.085, 180, 180 - 180 * val, width=0.025, facecolor=color, edgecolor="none", transform=ax01.transAxes))
        ax01.text(x, y - 0.01, f"{val*100:.1f}%", ha="center", va="center", fontsize=9, fontweight="bold", transform=ax01.transAxes)
        ax01.text(x, 0.28, label, ha="center", va="center", fontsize=7.8, transform=ax01.transAxes)
    ax01.text(0.5, 0.08, "33个代表性案例均具有机理-验证-预防与假设鉴别闭环", ha="center", fontsize=8.2, transform=ax01.transAxes)

    comp_labels = ["机理层", "传播步骤", "可观测量", "失效模式", "验证试验", "控制屏障", "诊断假设", "鉴别证据", "反事实测试", "判定规则"]
    comp_vals = [232, 232, 232, 174, 174, 174, 33, 99, 99, 66]
    comp_colors = ["#3A6EA5", "#4E9A80", "#6F9EAF", "#C05A5A", "#8C6BB1", "#C18F2C", "#999999", "#D4A34A", "#AA6F73", "#7C5E9E"]
    wedges, _ = ax10.pie(comp_vals, startangle=95, colors=comp_colors, wedgeprops={"width": 0.36, "edgecolor": "white"})
    ax10.text(0, 0.04, "1,521\n机理节点", ha="center", va="center", fontsize=11, fontweight="bold")
    ax10.set_title("机理增强图节点组成")
    legend_labels = [f"{n} {v}" for n, v in zip(comp_labels, comp_vals)]
    ax10.legend(wedges, legend_labels, loc="center left", bbox_to_anchor=(0.90, 0.5), fontsize=7, frameon=False)

    ax11.set_title("关系密度与结构摘要")
    rel_name_map = {
        "MECHANISM_PROPAGATES_TO": "机理传播",
        "HAS_MECHANISM_LAYER": "包含机理层",
        "MECHANISM_RESULTS_IN": "机理导致",
        "HAS_OBSERVABLE": "包含可观测量",
        "VALIDATES_ARCHETYPE": "验证机理原型",
        "HAS_PROPAGATION_STEP": "包含传播步骤",
        "HAS_FAILURE_MODE": "包含失效模式",
        "VERIFIED_BY_TEST": "由试验验证",
        "CONTROLLED_BY_BARRIER": "受屏障控制",
        "REQUIRES_DISCRIMINATING_EVIDENCE": "需要鉴别证据",
        "RESOLVED_BY_COUNTERFACTUAL_TEST": "由反事实测试消解",
        "HAS_DECISION_RULE": "包含判定规则",
    }
    rel_items = sorted(data["mechanism"]["relation_types"].items(), key=lambda item: item[1], reverse=True)[:8]
    y_pos = list(range(len(rel_items)))
    rel_vals = [v for _, v in rel_items]
    rel_names = [rel_name_map.get(k, k) for k, _ in rel_items]
    rel_colors = [NODE_COLORS[i % len(NODE_COLORS)][1] for i in range(len(rel_items))]
    ax11.barh(y_pos, rel_vals, color=rel_colors, alpha=0.78)
    ax11.set_yticks(y_pos, rel_names, fontsize=6.5)
    ax11.invert_yaxis()
    ax11.tick_params(axis="x", labelsize=7)
    ax11.grid(axis="x", linestyle=":", alpha=0.35)
    for i, v in enumerate(rel_vals):
        ax11.text(v + 4, i, str(v), va="center", fontsize=7)
    ax11.text(0.62, 0.11, "推理图 2,326节点 / 3,567关系\n机理子图 1,521节点 / 2,436关系\n成对竞争18例，单机理反事实15例", transform=ax11.transAxes, fontsize=8.2, va="bottom")
    fig.tight_layout(rect=[0, 0, 1, 0.92])
    return save_fig(fig, "图5_评估结果")


def fig_evidence_packet() -> Path:
    fig, ax = plt.subplots(figsize=(12.4, 7.0))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.96, "图6  代表性故障案例与普通RAG功能对比", ha="center", fontsize=14, fontweight="bold")
    cases = [
        ("变流器\nIGBT板", "电气热-绝缘", "粉尘/湿度\n驱动板短路\n高温跳闸", "更换驱动板\n通风除湿", "#C05A5A"),
        ("偏航\n液压压力", "液压流量受限", "151→27 bar\n300s恢复\n清洗后11s", "清洗柱塞孔\n调整缓冲开度", "#3A6EA5"),
        ("主轴承\n复合磨损", "接触疲劳磨损", "载荷/污染\n油膜破裂\n剥落", "停机换轴承\n密封滤油", "#C18F2C"),
        ("绕组温度\n跳变", "传感采集反馈", "70-135°C跳变\n停机稳定\n换通道排除PLC", "通道互换\n修复屏蔽接地", "#4E9A80"),
    ]
    ax.text(0.21, 0.86, "案例", ha="center", fontsize=10, fontweight="bold", color="#1F4E79")
    ax.text(0.43, 0.86, "机理类型", ha="center", fontsize=10, fontweight="bold", color="#7A4C12")
    ax.text(0.62, 0.86, "鉴别证据", ha="center", fontsize=10, fontweight="bold", color="#2B7A57")
    ax.text(0.80, 0.86, "现场动作", ha="center", fontsize=10, fontweight="bold", color="#7C5E9E")
    x_cols = [0.21, 0.43, 0.62, 0.80]
    for x in x_cols:
        ax.plot([x, x], [0.27, 0.82], color="#E6E6E6", linewidth=0.9)
    for i, (name, mech, evidence, action, color) in enumerate(cases):
        y = 0.76 - i * 0.13
        ax.plot([0.12, 0.89], [y, y], color="#EFEFEF", linewidth=1.0)
        draw_circle(ax, (0.13, y), 0.020, str(i + 1), fill=color, edge=color, size=8, bold=True)
        ax.text(0.21, y, name, ha="center", va="center", fontsize=8.2, fontweight="bold")
        draw_pill(ax, (0.35, y - 0.026), (0.16, 0.052), mech, fill="#FFF5E6", edge=color, size=7.3, bold=True)
        ax.text(0.62, y, evidence, ha="center", va="center", fontsize=7.1)
        ax.text(0.80, y, action, ha="center", va="center", fontsize=7.3, fontweight="bold")
        for x1, x2 in [(0.29, 0.35), (0.51, 0.56), (0.68, 0.73)]:
            arrow(ax, (x1, y), (x2, y))

    feature_labels = ["模型范围", "图谱路径", "机理链路", "假设鉴别", "多轮状态", "证据路径", "单步动作"]
    plain = [0.20, 0.35, 0.10, 0.00, 0.20, 0.45, 0.20]
    proposed = [1.00, 0.95, 1.00, 1.00, 0.90, 1.00, 1.00]
    ax.text(0.16, 0.19, "Plain RAG", ha="right", va="center", fontsize=8.5, fontweight="bold", color="#666666")
    ax.text(0.16, 0.11, "本方案", ha="right", va="center", fontsize=8.5, fontweight="bold", color="#23527C")
    for i, label in enumerate(feature_labels):
        x = 0.24 + i * 0.095
        ax.text(x, 0.235, label, ha="center", va="center", fontsize=7)
        accent = NODE_COLORS[i % len(NODE_COLORS)][1]
        for y, val, color in [(0.19, plain[i], "#999999"), (0.11, proposed[i], accent)]:
            ax.add_patch(Circle((x, y), 0.023, facecolor="#F2F2F2", edgecolor="#CCCCCC", linewidth=0.8))
            if val > 0:
                ax.add_patch(Wedge((x, y), 0.023, 90, 90 + 360 * val, facecolor=color, edgecolor="none", alpha=0.85))
            ax.add_patch(Circle((x, y), 0.012, facecolor="white", edgecolor="none"))
    draw_pill(ax, (0.22, 0.018), (0.56, 0.050), "输出合同：故障含义 | 适用范围 | 可能机理 | 一个验证动作 | 接受准则 | 证据路径", fill="#F8F0FA", edge="#7C5E9E", size=8.0, bold=True)
    ax.text(0.5, 0.275, "上部为论文四个代表性案例的机理-证据-动作矩阵；下部为普通RAG与证据约束策略链的功能覆盖差异。", ha="center", fontsize=8.4)
    return save_fig(fig, "图6_案例与对比")


def add_info_table(doc: Document) -> None:
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("交底书名称", "一种面向风电机组运维的本地化证据约束大语言模型策略链故障诊断方法及系统"),
        ("技术联系人姓名及其电话、email", "[待填写：技术联系人、电话、邮箱]"),
        ("经办人姓名及其电话、email", "[待填写：经办人、电话、邮箱；如不填写，则默认技术联系人为负责人]"),
        ("是否仅用于申请政府高新资质或政府项目", "[待确认]"),
        ("是否将来申请国外专利", "[待确认]"),
        ("备注", "本交底书根据论文《Evidence-Grounded LLM Strategy Chains for Localized Wind Turbine Fault Operations and Maintenance》整理，申请主体、发明人、公开状态和提交策略需由申请人进一步确认。"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        set_cell_text(row.cells[0], k, bold=True)
        set_cell_shading(row.cells[0], "EDEDED")
        set_cell_text(row.cells[1], v)


def add_case_table(doc: Document) -> None:
    add_h2(doc, "典型实施例中的故障处置样例")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    heads = ["案例", "机理链", "关键证据", "建议动作"]
    for c, h in zip(table.rows[0].cells, heads):
        set_cell_text(c, h, bold=True)
        set_cell_shading(c, "EDEDED")
    rows = [
        [
            "变流器IGBT板故障",
            "电网波动、粉尘、盐雾、湿度或散热不良导致驱动板短路、电容损坏、温升和变流器跳闸。",
            "ISU SC、INU SC报警；K8接触器不能闭合；温度传感器未显示异常；改造后运行温度降低。",
            "更换IGBT驱动板，清洁过滤网，改善通风除湿，并监测板卡温度。",
        ],
        [
            "偏航液压压力异常",
            "隐藏阻尼缓冲结构堵塞限制液压回路油流，造成压力恢复慢和液压站欠压。",
            "手动释放偏航制动时压力由151 bar降至27 bar，恢复至150 bar约300 s；清洗后11 s升至153 bar。",
            "清洗柱塞孔和阻尼缓冲件，调整缓冲开度，必要时更换改进缓冲件并更新液压图纸。",
        ],
        [
            "主轴承复合磨损",
            "高海拔温差、湍流载荷、润滑污染、同轴度偏差和密封老化导致油膜破裂与剥落。",
            "振动速度4.8 mm/s，峭度8.7，金属碎屑1200 ppm，同轴度偏差0.15 mm，内外圈剥落。",
            "停机，更换轴承，清洗润滑腔，修复轴颈，升级密封及在线滤油。",
        ],
        [
            "发电机绕组温度跳变",
            "屏蔽接地不良、线缆布置不规范和电磁耦合导致Pt100模拟温度信号畸变。",
            "温度在70至135摄氏度之间跳变；停机时稳定；并网后波动增强；通道互换排除PLC模块故障。",
            "检查Pt100阻值，互换采集通道，修复屏蔽接地，更换信号线，必要时切换备用传感器。",
        ],
    ]
    for data in rows:
        cells = table.add_row().cells
        for c, v in zip(cells, data):
            set_cell_text(c, v)


def build_doc() -> Path:
    doc = Document()
    setup_doc(doc)
    add_title(doc, "专利申请技术交底书")
    add_info_table(doc)

    add_h1(doc, "一、与本专利最接近的现有技术")
    add_h2(doc, "1、现有技术的方案简述")
    add_para(doc, "风电机组运维现场通常依赖故障码表、设备手册、历史维修记录、工单、案例报告以及技术人员经验进行故障判断。现有数字化运维方案主要包括关键字检索、SCADA监测平台、静态知识库、预测性维护模型以及普通大语言模型问答或普通RAG检索问答。上述方案能够在一定程度上提高资料检索效率，但多数方案仍以非结构化文本片段、单一故障码或单一监测信号为处理对象。")
    add_para(doc, "在普通检索或普通RAG方案中，系统通常根据用户输入召回若干文本片段，再由语言模型生成自然语言答复。该过程缺少对风场、机型、品牌、控制系统、部件和会话状态的严格绑定，也缺少对“故障码—适用机型—系统—部件—原因—处置—来源文档”之间关系的显式表达。对于同一数字故障码在不同品牌或不同机型中含义不同的情况，普通文本检索容易返回看似相关但不适用于当前机组的内容。")
    add_para(doc, "现有基于传感器或SCADA的故障诊断方法侧重数值异常检测、振动分析或预测模型，能够发现状态异常，但通常不能直接解决文档级维修推理问题，例如定位适用手册、解释型号限定的故障码、给出可审计的来源路径以及在多轮现场反馈中继续执行同一故障分支。")
    add_h2(doc, "2、现有技术的客观缺点")
    add_bullets(doc, [
        "缺少型号和场景范围约束：同一故障码在不同机型、品牌或控制系统下可能对应不同含义，普通文本检索不能可靠消歧。",
        "缺少结构化证据路径：普通RAG返回片段，但不显式保存故障、部件、原因、动作、来源文档及机理之间的关系，答案难以审计。",
        "缺少机理级推理链：传统故障-现象-处置模式不能解释为何检查某个信号、如何验证根因以及如何预防复发。",
        "缺少假设鉴别能力：现场多个原因可能产生相同报警，现有方案难以区分真实物理故障、传感误差、边界条件影响或控制时序伪原因。",
        "缺少多轮状态保持：维修人员给出简短反馈时，普通问答系统往往将其视为新问题，不能延续前一轮验证动作和分支判断。",
    ])

    add_h1(doc, "二、本专利的技术")
    add_h2(doc, "（1）具体方案")
    add_para(doc, "本专利提出一种面向风电机组运维的本地化证据约束大语言模型策略链故障诊断方法及系统。该方案将本地运维文档转换为结构化故障条目、知识图谱、机理模板和可读知识页面，在问答阶段利用本地证据包约束大语言模型，使其输出具有适用范围、可能机理、现场验证动作、接受准则和来源路径的运维处置建议。")
    for img, cap in [
        (fig_main_flow(), "图1  本地化风电运维证据约束LLM系统架构"),
        (fig_architecture(), "图2  本地运维知识构建流程"),
        (fig_scope_strategy(), "图3  模型范围绑定、证据包生成与策略链交互"),
        (fig_mechanism(), "图4  机理增强推理图与假设鉴别闭环"),
        (fig_metrics(), "图5  知识资产规模、机理图谱组成与覆盖效果"),
        (fig_evidence_packet(), "图6  代表性故障案例与普通RAG功能对比"),
    ]:
        doc.add_picture(str(img), width=Inches(6.6))
        add_caption(doc, cap)

    add_para(doc, "如图1所示，系统包括数据层、知识层和推理层。数据层接入本地手册、故障码表、案例报告、维修工单和问答轨迹；知识层执行字段归一化、证据化故障条目构建、知识图谱构建、机理模板生成以及Wiki与索引输出；推理层执行意图路由、范围绑定、图增强检索、策略链控制和受约束生成。大语言模型不作为独立事实来源，而是仅对证据包和当前故障状态进行语言组织和策略控制。")
    add_para(doc, "如图2所示，知识构建流程包括文档登记、文本与表格解析、领域字段识别、别名归一化、证据化故障条目生成以及关系与机理生成。该流程的输出包括故障索引、知识图谱、机理模板和可读Wiki页面，从而使非结构化本地资料变成可检索、可回溯、可供策略链调用的结构化证据资产。")
    add_para(doc, "故障条目可表示为式（1）。")
    add_math(doc, r"f_i=(c_i,m_i,b_i,s_i,p_i,r_i,a_i,e_i)", 1)
    add_para(doc, "其中，c_i表示故障码或报警名称，m_i表示适用机型，b_i表示品牌或控制器族，s_i表示系统，p_i表示部件，r_i表示原因描述，a_i表示处置或复位规则，e_i表示证据指针。证据指针至少包括来源文件、路径、章节或页码提示以及原始抽取文本。")
    add_para(doc, "机理模板可表示为式（2）。")
    add_math(doc, r"g_j=(u_j,v_j,q_j,k_j)", 2)
    add_para(doc, "其中，u_j表示典型症状，v_j表示因果因素，q_j表示优先检查项，k_j表示接受准则或分支条件。系统根据系统类型、部件、故障表现、原因术语和推荐动作，将多个故障条目映射至同一机理模板，从而避免仅以孤立故障码输出处置建议。")
    add_para(doc, "范围绑定得分可表示为式（3）。")
    add_math(doc, r"score(f_i,q)=\lambda_c C_i+\lambda_m M_i+\lambda_b B_i+\lambda_s S_i+\lambda_h H_i+\lambda_e E_i", 3)
    add_para(doc, "其中，C_i表示故障码匹配信号，M_i表示机型一致性，B_i表示品牌一致性，S_i表示系统或部件术语重合度，H_i表示与前序会话状态的一致性，E_i表示来源证据强度。当多个候选故障条目的机型范围互不兼容时，系统不直接合并输出，而是返回缺失范围条件或分别列出候选范围。如图3所示，范围绑定结果与结构化故障条目、图谱邻居、机理链路和来源路径共同形成证据包，策略链据此输出故障含义、可能机理、一个现场验证动作、接受准则和证据路径。")
    add_para(doc, "如图4所示，机理增强推理图包括故障案例、机理原型、机理层、传播步骤、失效模式、可观测量、验证试验、控制屏障、诊断假设、鉴别证据、反事实测试和判定规则。对于每一故障案例，系统形成式（4）所示的闭环链路。")
    add_math(doc, r"x\rightarrow a\rightarrow l\rightarrow p\rightarrow z\rightarrow o\rightarrow t\rightarrow b", 4)
    add_para(doc, "其中，x为故障案例，a为机理原型，l为机理层，p为传播步骤，z为失效模式，o为可观测量，t为验证试验，b为控制屏障。")
    add_para(doc, "系统还为诊断鉴别建立式（5）所示的假设链路。")
    add_math(doc, r"x\rightarrow h\rightarrow \{e_d,t_c,r_d\}", 5)
    add_para(doc, "其中，h为诊断假设，e_d为鉴别证据，t_c为反事实测试，r_d为判定规则。当故障案例匹配两个机理原型时，系统形成成对竞争假设；当故障案例只有单一主导机理时，系统形成单机理反事实假设，用于排除测量伪差、边界条件影响或控制时序伪原因。")
    add_para(doc, "大语言模型策略链维护当前故障对象、模型范围、已召回证据、已完成验证动作、用户反馈和剩余分支。每轮输出限定为故障含义、适用范围、可能机理、一个优先现场验证动作、接受准则和证据路径。维修人员返回简短反馈后，系统将该反馈解释为上一验证动作的观测结果，并根据判定规则进入下一分支。")

    add_case_table(doc)

    add_h2(doc, "（2）技术效果")
    add_para(doc, "通过将本地运维文档转换为结构化故障条目和知识图谱，本方案能够在故障码检索时同时考虑机型、品牌、系统、部件和会话状态，减少跨机型误匹配导致的错误建议。")
    add_para(doc, "通过证据包约束大语言模型，本方案使生成答案中的关键结论均可回溯至本地文档或图谱路径，便于维修人员、技术专家和审计人员复核。")
    add_para(doc, "通过机理增强推理链，本方案把传统故障-证据-处置关系扩展为故障-机理原型-失效模式-传播步骤-可观测量-验证试验-控制屏障的闭环链路，使系统不仅能够给出处置动作，还能解释为何执行该检查以及如何预防复发。")
    add_para(doc, "通过诊断假设和反事实测试，本方案能够在多个可能原因之间形成可验证的鉴别路径，避免将传感器误差、采集链路问题、边界条件或控制时序伪原因误判为真实物理故障。")
    add_para(doc, "根据论文实施数据，如图5所示，原型系统覆盖11865条原始故障记录、22680个图谱节点、79474条图谱关系、4849个故障码节点和8295个来源文档节点；机理增强推理图包含2326个节点和3567条关系，在33个代表性故障案例上机理闭环覆盖率、假设鉴别覆盖率和机理画像完整率均达到100%。如图6所示，代表性案例覆盖变流器IGBT板、偏航液压压力、主轴承复合磨损和绕组温度跳变等不同机理类型，相比普通RAG更强调模型范围、机理链路、证据路径和多轮现场状态。")

    add_h1(doc, "三、上述技术方案是否有替代方案")
    add_bullets(doc, [
        "文档来源可替换为数据库、对象存储、企业知识库、文件系统同步目录、SCADA事件记录或维修工单系统，只要能够保留来源路径和原始证据片段即可。",
        "知识图谱可采用文件型图数据、关系型数据库、图数据库、向量数据库与结构化索引混合实现，图谱存储方式不限制本方案的保护范围。",
        "大语言模型可为本地部署模型、企业内网模型服务或经批准的云端模型接口；事实生成仍由本地证据包约束。",
        "范围绑定得分中的权重λ_c、λ_m、λ_b、λ_s、λ_h、λ_e可通过人工配置、规则学习或历史维修反馈自动调整。",
        "机理模板可由人工专家配置、从历史案例半自动抽取或由规则与模型共同生成；机理原型数量和名称可根据风机类型扩展。",
        "策略链可输出单一优先动作，也可输出按风险等级排序的多个动作；对于高风险操作可增加人工确认节点。",
        "诊断假设可采用成对竞争、单机理反事实、多假设排序或概率图模型等实现形式。",
    ])

    add_h1(doc, "四、本发明的关键技术点")
    add_numbered(doc, [
        "将本地运维文档解析为带来源证据指针的结构化故障条目，而不是仅建立普通文本片段索引。",
        "在故障码问答前进行风场、机型、品牌、系统、部件和会话状态的范围绑定，解决同码异义和跨机型误匹配问题。",
        "构建包含机理原型、失效模式、传播步骤、可观测量、验证试验和控制屏障的机理增强推理图。",
        "为每个故障案例建立诊断假设、鉴别证据、反事实测试和判定规则，以支持可证伪的现场排查。",
        "将结构化故障条目、图谱邻居、机理链和来源路径压缩为证据包，并用证据包约束大语言模型生成。",
        "维护多轮故障策略状态，使维修人员的简短反馈能够被解释为上一现场验证动作的观测结果并进入下一分支。",
        "输出包含故障含义、适用范围、可能机理、优先验证动作、接受准则和证据路径的运维处置建议。",
    ])

    add_h1(doc, "五、其他有助于理解本申请提案的技术资料")
    add_bullets(doc, [
        "论文题目：Evidence-Grounded LLM Strategy Chains for Localized Wind Turbine Fault Operations and Maintenance。",
        "本地论文源文件：bare_jrnl_new_sample4.tex；论文图文件目录：figures/。",
        "实施规模：11865条原始故障记录、22680个图谱节点、79474条图谱关系、4849个故障码节点、8295个来源文档节点、66种机型、29个风场和12个机理模板。",
        "机理增强评估：33个代表性故障案例均具备机理原型、失效模式、可观测量、验证试验、控制屏障、诊断假设、鉴别证据、反事实测试和判定规则。",
        "待申请人确认事项：申请人、发明人、技术联系人、经办人、是否已公开论文或演示、是否计划国外申请、是否包含未公开的企业文档或客户数据。",
    ])

    path = OUT / "风电运维证据约束LLM策略链_专利技术交底书.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    result = build_doc()
    print(result)
